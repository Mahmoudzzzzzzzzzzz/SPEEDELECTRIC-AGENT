#!/usr/bin/env python3
import requests
import json
import os
import time
import io
import pandas as pd
from docx import Document
import unittest
from datetime import datetime, timedelta
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Get backend URL from frontend .env file
def get_backend_url():
    with open('/app/frontend/.env', 'r') as f:
        for line in f:
            if line.startswith('REACT_APP_BACKEND_URL='):
                return line.strip().split('=')[1].strip('"\'')
    raise ValueError("Backend URL not found in frontend/.env")

# Base URL for API requests
BASE_URL = f"{get_backend_url()}/api"
logger.info(f"Using backend URL: {BASE_URL}")

# Helper function to create test files
def create_test_docx():
    """Create a test Word document with customer data"""
    doc = Document()
    doc.add_heading('Customer Data', 0)
    
    # Add some customers
    doc.add_paragraph('Customer 1:')
    doc.add_paragraph('Name: John Smith')
    doc.add_paragraph('Email: john.smith@example.com')
    doc.add_paragraph('Company: ABC Corp')
    doc.add_paragraph('Phone: 555-123-4567')
    doc.add_paragraph('Address: 123 Main St, Anytown, USA')
    
    doc.add_paragraph('Customer 2:')
    doc.add_paragraph('Name: Jane Doe')
    doc.add_paragraph('Email: jane.doe@example.com')
    doc.add_paragraph('Company: XYZ Inc')
    doc.add_paragraph('Phone: 555-987-6543')
    doc.add_paragraph('Address: 456 Oak Ave, Somewhere, USA')
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_test_excel():
    """Create a test Excel file with customer data"""
    data = {
        'name': ['Alice Johnson', 'Bob Williams', 'Carol Brown'],
        'email': ['alice@example.com', 'bob@example.com', 'carol@example.com'],
        'company': ['Tech Solutions', 'Marketing Pro', 'Design Studio'],
        'phone': ['555-111-2222', '555-333-4444', '555-555-6666'],
        'address': ['789 Pine St, Somewhere, USA', '101 Maple Dr, Anywhere, USA', '202 Cedar Ln, Nowhere, USA']
    }
    
    df = pd.DataFrame(data)
    
    # Save to memory
    file_stream = io.BytesIO()
    df.to_excel(file_stream, index=False, engine='openpyxl')
    file_stream.seek(0)
    return file_stream

class BidTrackerAPITest(unittest.TestCase):
    """Test suite for Bid Tracker API"""
    
    def setUp(self):
        """Set up test data"""
        self.customer_ids = []
        self.template_ids = []
        self.campaign_ids = []
        self.followup_ids = []
    
    def test_01_api_root(self):
        """Test API root endpoint"""
        logger.info("Testing API root endpoint")
        response = requests.get(f"{BASE_URL}/")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["message"], "Bid Tracker API")
        self.assertEqual(data["version"], "1.0.0")
        logger.info("API root endpoint test passed")
    
    def test_02_create_customer(self):
        """Test creating a customer"""
        logger.info("Testing customer creation")
        
        # Create test customer
        customer_data = {
            "name": "Test Customer",
            "email": "test.customer@example.com",
            "company": "Test Company",
            "phone": "555-TEST",
            "address": "Test Address",
            "tags": ["test", "api"],
            "notes": "Created for API testing"
        }
        
        response = requests.post(f"{BASE_URL}/customers", json=customer_data)
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["name"], customer_data["name"])
        self.assertEqual(data["email"], customer_data["email"])
        self.assertEqual(data["company"], customer_data["company"])
        self.assertEqual(data["tags"], customer_data["tags"])
        
        # Save customer ID for later tests
        self.customer_ids.append(data["id"])
        logger.info(f"Created customer with ID: {data['id']}")
        
        # Create a second customer for testing
        customer_data2 = {
            "name": "Another Customer",
            "email": "another.customer@example.com",
            "company": "Another Company",
            "phone": "555-ANOTHER",
            "address": "Another Address",
            "tags": ["test", "another"],
            "notes": "Another test customer"
        }
        
        response = requests.post(f"{BASE_URL}/customers", json=customer_data2)
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.customer_ids.append(data["id"])
        logger.info(f"Created second customer with ID: {data['id']}")
    
    def test_03_get_customers(self):
        """Test retrieving customers"""
        logger.info("Testing customer retrieval")
        
        # Get all customers
        response = requests.get(f"{BASE_URL}/customers")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertIsInstance(data, list)
        self.assertGreaterEqual(len(data), len(self.customer_ids))
        
        # Test pagination
        response = requests.get(f"{BASE_URL}/customers?skip=0&limit=1")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(len(data), 1)
        
        # Test filtering by status
        response = requests.get(f"{BASE_URL}/customers?status=active")
        self.assertEqual(response.status_code, 200)
        
        logger.info("Customer retrieval test passed")
    
    def test_04_get_customer_by_id(self):
        """Test retrieving a specific customer"""
        if not self.customer_ids:
            self.skipTest("No customer IDs available")
        
        logger.info(f"Testing retrieval of customer with ID: {self.customer_ids[0]}")
        
        response = requests.get(f"{BASE_URL}/customers/{self.customer_ids[0]}")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["id"], self.customer_ids[0])
        
        # Test with non-existent ID
        response = requests.get(f"{BASE_URL}/customers/nonexistent-id")
        self.assertEqual(response.status_code, 404)
        
        logger.info("Customer retrieval by ID test passed")
    
    def test_05_update_customer(self):
        """Test updating a customer"""
        if not self.customer_ids:
            self.skipTest("No customer IDs available")
        
        logger.info(f"Testing update of customer with ID: {self.customer_ids[0]}")
        
        update_data = {
            "name": "Updated Customer Name",
            "notes": "Updated via API test",
            "status": "prospect"
        }
        
        response = requests.put(f"{BASE_URL}/customers/{self.customer_ids[0]}", json=update_data)
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["name"], update_data["name"])
        self.assertEqual(data["notes"], update_data["notes"])
        self.assertEqual(data["status"], update_data["status"])
        
        # Verify update with a GET request
        response = requests.get(f"{BASE_URL}/customers/{self.customer_ids[0]}")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["name"], update_data["name"])
        
        logger.info("Customer update test passed")
    
    def test_06_import_customers_docx(self):
        """Test importing customers from Word document"""
        logger.info("Testing customer import from Word document")
        
        # Create test Word document
        file_stream = create_test_docx()
        
        # Upload file
        files = {'file': ('test_customers.docx', file_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(f"{BASE_URL}/customers/import", files=files)
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        
        self.assertIn("message", data)
        self.assertIn("imported_count", data)
        self.assertIn("customers", data)
        self.assertGreater(data["imported_count"], 0)
        
        # Save customer IDs for later tests
        for customer in data["customers"]:
            self.customer_ids.append(customer["id"])
        
        logger.info(f"Imported {data['imported_count']} customers from Word document")
    
    def test_07_import_customers_excel(self):
        """Test importing customers from Excel file"""
        logger.info("Testing customer import from Excel file")
        
        # Create test Excel file
        file_stream = create_test_excel()
        
        # Upload file
        files = {'file': ('test_customers.xlsx', file_stream, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
        response = requests.post(f"{BASE_URL}/customers/import", files=files)
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        
        self.assertIn("message", data)
        self.assertIn("imported_count", data)
        self.assertIn("customers", data)
        self.assertGreater(data["imported_count"], 0)
        
        # Save customer IDs for later tests
        for customer in data["customers"]:
            self.customer_ids.append(customer["id"])
        
        logger.info(f"Imported {data['imported_count']} customers from Excel file")
    
    def test_08_create_email_template(self):
        """Test creating email templates"""
        logger.info("Testing email template creation")
        
        # Create proposal template
        proposal_template = {
            "name": "Test Proposal Template",
            "subject": "Proposal for {{company}}",
            "body": "Dear {{name}},\n\nWe are pleased to submit our proposal for {{company}}.\n\nRegards,\nBid Tracker Team",
            "template_type": "proposal",
            "variables": ["name", "company"]
        }
        
        response = requests.post(f"{BASE_URL}/templates", json=proposal_template)
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.template_ids.append(data["id"])
        logger.info(f"Created proposal template with ID: {data['id']}")
        
        # Create follow-up template
        followup_template = {
            "name": "Test Follow-up Template",
            "subject": "Follow-up on our proposal",
            "body": "Dear {{name}},\n\nI wanted to follow up on the proposal we sent for {{company}}.\n\nRegards,\nBid Tracker Team",
            "template_type": "follow_up",
            "variables": ["name", "company"]
        }
        
        response = requests.post(f"{BASE_URL}/templates", json=followup_template)
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.template_ids.append(data["id"])
        logger.info(f"Created follow-up template with ID: {data['id']}")
        
        # Create general template
        general_template = {
            "name": "Test General Template",
            "subject": "General information",
            "body": "Dear {{name}},\n\nThis is a general message for {{company}}.\n\nRegards,\nBid Tracker Team",
            "template_type": "general",
            "variables": ["name", "company"]
        }
        
        response = requests.post(f"{BASE_URL}/templates", json=general_template)
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.template_ids.append(data["id"])
        logger.info(f"Created general template with ID: {data['id']}")
    
    def test_09_get_templates(self):
        """Test retrieving email templates"""
        logger.info("Testing email template retrieval")
        
        # Get all templates
        response = requests.get(f"{BASE_URL}/templates")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertIsInstance(data, list)
        self.assertGreaterEqual(len(data), len(self.template_ids))
        
        # Test filtering by template type
        response = requests.get(f"{BASE_URL}/templates?template_type=proposal")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertGreater(len(data), 0)
        for template in data:
            self.assertEqual(template["template_type"], "proposal")
        
        logger.info("Email template retrieval test passed")
    
    def test_10_get_template_by_id(self):
        """Test retrieving a specific template"""
        if not self.template_ids:
            self.skipTest("No template IDs available")
        
        logger.info(f"Testing retrieval of template with ID: {self.template_ids[0]}")
        
        response = requests.get(f"{BASE_URL}/templates/{self.template_ids[0]}")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["id"], self.template_ids[0])
        
        # Test with non-existent ID
        response = requests.get(f"{BASE_URL}/templates/nonexistent-id")
        self.assertEqual(response.status_code, 404)
        
        logger.info("Template retrieval by ID test passed")
    
    def test_11_update_template(self):
        """Test updating an email template"""
        if not self.template_ids:
            self.skipTest("No template IDs available")
        
        logger.info(f"Testing update of template with ID: {self.template_ids[0]}")
        
        update_data = {
            "name": "Updated Template Name",
            "subject": "Updated Subject for {{company}}",
            "body": "Updated body with {{name}} and {{company}} variables",
            "template_type": "proposal",
            "variables": ["name", "company"]
        }
        
        response = requests.put(f"{BASE_URL}/templates/{self.template_ids[0]}", json=update_data)
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["name"], update_data["name"])
        self.assertEqual(data["subject"], update_data["subject"])
        self.assertEqual(data["body"], update_data["body"])
        
        logger.info("Template update test passed")
    
    def test_12_create_campaign(self):
        """Test creating a campaign"""
        if not self.template_ids or not self.customer_ids:
            self.skipTest("No template or customer IDs available")
        
        logger.info("Testing campaign creation")
        
        # Create campaign
        campaign_data = {
            "name": "Test Campaign",
            "template_id": self.template_ids[0],
            "customer_ids": self.customer_ids[:2],  # Use first two customers
            "scheduled_at": (datetime.utcnow() + timedelta(days=1)).isoformat()
        }
        
        response = requests.post(f"{BASE_URL}/campaigns", json=campaign_data)
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["name"], campaign_data["name"])
        self.assertEqual(data["template_id"], campaign_data["template_id"])
        self.assertEqual(data["customer_ids"], campaign_data["customer_ids"])
        
        # Save campaign ID for later tests
        self.campaign_ids.append(data["id"])
        logger.info(f"Created campaign with ID: {data['id']}")
    
    def test_13_get_campaigns(self):
        """Test retrieving campaigns"""
        logger.info("Testing campaign retrieval")
        
        # Get all campaigns
        response = requests.get(f"{BASE_URL}/campaigns")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertIsInstance(data, list)
        self.assertGreaterEqual(len(data), len(self.campaign_ids))
        
        logger.info("Campaign retrieval test passed")
    
    def test_14_get_campaign_by_id(self):
        """Test retrieving a specific campaign"""
        if not self.campaign_ids:
            self.skipTest("No campaign IDs available")
        
        logger.info(f"Testing retrieval of campaign with ID: {self.campaign_ids[0]}")
        
        response = requests.get(f"{BASE_URL}/campaigns/{self.campaign_ids[0]}")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["id"], self.campaign_ids[0])
        
        # Test with non-existent ID
        response = requests.get(f"{BASE_URL}/campaigns/nonexistent-id")
        self.assertEqual(response.status_code, 404)
        
        logger.info("Campaign retrieval by ID test passed")
    
    def test_15_create_followup(self):
        """Test creating a follow-up"""
        if not self.template_ids or not self.customer_ids:
            self.skipTest("No template or customer IDs available")
        
        logger.info("Testing follow-up creation")
        
        # Create follow-up
        followup_data = {
            "customer_id": self.customer_ids[0],
            "template_id": self.template_ids[1],  # Use follow-up template
            "due_date": (datetime.utcnow() + timedelta(days=3)).isoformat(),
            "notes": "Follow up on proposal"
        }
        
        response = requests.post(f"{BASE_URL}/followups", json=followup_data)
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertEqual(data["customer_id"], followup_data["customer_id"])
        self.assertEqual(data["template_id"], followup_data["template_id"])
        self.assertEqual(data["notes"], followup_data["notes"])
        
        # Save follow-up ID for later tests
        self.followup_ids.append(data["id"])
        logger.info(f"Created follow-up with ID: {data['id']}")
    
    def test_16_get_followups(self):
        """Test retrieving follow-ups"""
        logger.info("Testing follow-up retrieval")
        
        # Get all follow-ups
        response = requests.get(f"{BASE_URL}/followups")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertIsInstance(data, list)
        self.assertGreaterEqual(len(data), len(self.followup_ids))
        
        # Test filtering by status
        response = requests.get(f"{BASE_URL}/followups?status=pending")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        for followup in data:
            self.assertEqual(followup["status"], "pending")
        
        # Test due soon filter
        response = requests.get(f"{BASE_URL}/followups?due_soon=true")
        self.assertEqual(response.status_code, 200)
        
        logger.info("Follow-up retrieval test passed")
    
    def test_17_dashboard_stats(self):
        """Test dashboard statistics"""
        logger.info("Testing dashboard statistics")
        
        response = requests.get(f"{BASE_URL}/dashboard/stats")
        self.assertEqual(response.status_code, 200)
        
        data = response.json()
        self.assertIn("customers", data)
        self.assertIn("campaigns", data)
        self.assertIn("followups", data)
        self.assertIn("templates", data)
        
        # Verify customer count
        self.assertGreaterEqual(data["customers"]["total"], len(self.customer_ids))
        
        # Verify template count
        self.assertGreaterEqual(data["templates"]["total"], len(self.template_ids))
        
        # Verify campaign count
        self.assertGreaterEqual(data["campaigns"]["total"], len(self.campaign_ids))
        
        logger.info("Dashboard statistics test passed")
    
    def test_18_delete_template(self):
        """Test deleting an email template"""
        if not self.template_ids:
            self.skipTest("No template IDs available")
        
        logger.info(f"Testing deletion of template with ID: {self.template_ids[-1]}")
        
        response = requests.delete(f"{BASE_URL}/templates/{self.template_ids[-1]}")
        self.assertEqual(response.status_code, 200)
        
        # Verify deletion
        response = requests.get(f"{BASE_URL}/templates/{self.template_ids[-1]}")
        self.assertEqual(response.status_code, 404)
        
        logger.info("Template deletion test passed")
    
    def test_19_delete_customer(self):
        """Test deleting a customer"""
        if not self.customer_ids:
            self.skipTest("No customer IDs available")
        
        logger.info(f"Testing deletion of customer with ID: {self.customer_ids[-1]}")
        
        response = requests.delete(f"{BASE_URL}/customers/{self.customer_ids[-1]}")
        self.assertEqual(response.status_code, 200)
        
        # Verify deletion
        response = requests.get(f"{BASE_URL}/customers/{self.customer_ids[-1]}")
        self.assertEqual(response.status_code, 404)
        
        logger.info("Customer deletion test passed")
    
    def test_20_error_handling(self):
        """Test error handling for invalid data"""
        logger.info("Testing error handling")
        
        # Test invalid customer creation (missing required field)
        invalid_customer = {
            "company": "Invalid Customer"
            # Missing name and email
        }
        
        response = requests.post(f"{BASE_URL}/customers", json=invalid_customer)
        self.assertNotEqual(response.status_code, 200)
        
        # Test invalid template creation
        invalid_template = {
            "name": "Invalid Template"
            # Missing subject and body
        }
        
        response = requests.post(f"{BASE_URL}/templates", json=invalid_template)
        self.assertNotEqual(response.status_code, 200)
        
        # Test invalid file upload
        response = requests.post(f"{BASE_URL}/customers/import")
        self.assertNotEqual(response.status_code, 200)
        
        logger.info("Error handling test passed")


if __name__ == "__main__":
    # Run the tests
    unittest.main(argv=['first-arg-is-ignored'], exit=False)